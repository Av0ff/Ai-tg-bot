import os
from dataclasses import dataclass
from typing import Iterable, List

from docx import Document as DocxDocument
from pypdf import PdfReader


@dataclass(frozen=True)
class ParsedDocument:
    source: str
    text: str


def load_documents_from_dir(path: str) -> List[ParsedDocument]:
    docs: List[ParsedDocument] = []
    for entry in _iter_files(path):
        text = _read_file(entry)
        if text.strip():
            docs.append(ParsedDocument(source=os.path.basename(entry), text=text))
    return docs


def _iter_files(path: str) -> Iterable[str]:
    if not os.path.isdir(path):
        return []
    for name in os.listdir(path):
        full_path = os.path.join(path, name)
        if os.path.isdir(full_path):
            continue
        ext = os.path.splitext(name)[1].lower()
        if ext in {".pdf", ".docx", ".txt"}:
            yield full_path


def _read_file(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _read_pdf(path)
    if ext == ".docx":
        return _read_docx(path)
    if ext == ".txt":
        return _read_txt(path)
    return ""


def _read_pdf(path: str) -> str:
    reader = PdfReader(path)
    parts: List[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text:
            parts.append(text)
    return "\n".join(parts)


def _read_docx(path: str) -> str:
    doc = DocxDocument(path)
    parts = [para.text for para in doc.paragraphs if para.text]
    return "\n".join(parts)


def _read_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8") as handle:
        return handle.read()
